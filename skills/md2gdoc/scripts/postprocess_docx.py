"""
Post-process a pandoc-generated .docx to fix formatting issues
that Google Docs import doesn't handle well:
  1. Table borders on all cells
  2. Header row shading (gray background)
  3. Code block background (light gray) + monospace font
  4. Bold header row text
  5. Paragraph spacing (body text, lists, code blocks)
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def add_cell_borders(cell, color="000000", size="4"):
    """Add solid borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing borders if any
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)

    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_cell_shading(cell, color="D9D9D9"):
    """Add background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)

    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    tcPr.append(shading)


def set_cell_text_bold(cell):
    """Make all text in a cell bold."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def set_cell_font(cell, font_name="Arial", font_size=Pt(10)):
    """Set font for all runs in a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size


def is_code_block(paragraph):
    """Check if a paragraph is a code block (Source Code style or similar)."""
    style_name = paragraph.style.name if paragraph.style else ""
    code_indicators = [
        'Source Code', 'Verbatim', 'Code', 'source-code',
        'verbatim-char', 'VerbatimChar'
    ]
    # Check style name
    for indicator in code_indicators:
        if indicator.lower() in style_name.lower():
            return True

    # Check if all runs use monospace font
    if paragraph.runs:
        mono_fonts = {'Courier New', 'Consolas', 'Roboto Mono', 'Source Code Pro', 'monospace'}
        all_mono = all(
            run.font.name in mono_fonts
            for run in paragraph.runs
            if run.font.name is not None
        )
        if all_mono and len(paragraph.runs) > 0:
            # Check that at least one run has a known mono font
            has_mono = any(
                run.font.name in mono_fonts
                for run in paragraph.runs
                if run.font.name is not None
            )
            if has_mono:
                return True

    return False


def style_code_block(paragraph, bg_color="F3F3F3", font_name="Roboto Mono", font_size=Pt(9)):
    """Apply code block styling: background shading + monospace font."""
    # Add paragraph-level shading
    pPr = paragraph._p.get_or_add_pPr()

    existing_shd = pPr.find(qn('w:shd'))
    if existing_shd is not None:
        pPr.remove(existing_shd)

    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>'
    )
    pPr.append(shading)

    # Add left/right indentation for visual padding
    existing_ind = pPr.find(qn('w:ind'))
    if existing_ind is None:
        ind = parse_xml(
            f'<w:ind {nsdecls("w")} w:left="288" w:right="288"/>'
        )
        pPr.append(ind)

    # Set monospace font on all runs
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)  # Dark gray like reference


def set_table_full_width(table):
    """Set table width to 100% of the page width with auto-fit columns."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove existing width setting
    existing_w = tblPr.find(qn('w:tblW'))
    if existing_w is not None:
        tblPr.remove(existing_w)

    # Set to 100% width (5000 = 100% in fiftieths of a percent)
    tbl_width = parse_xml(
        f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>'
    )
    tblPr.append(tbl_width)

    # Set autofit layout so columns redistribute to fill the full width
    existing_layout = tblPr.find(qn('w:tblLayout'))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    layout = parse_xml(
        f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>'
    )
    tblPr.append(layout)

    # Remove fixed column widths from the grid so autofit can work
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn('w:gridCol')):
            if gridCol.get(qn('w:w')) is not None:
                gridCol.attrib.pop(qn('w:w'))

    # Remove fixed widths from individual cells
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)


def fix_tables(doc):
    """Add borders, header styling, and full-width to all tables."""
    for table in doc.tables:
        # Stretch table to full page width
        set_table_full_width(table)

        # Process all rows
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                add_cell_borders(cell)
                set_cell_font(cell, "Arial", Pt(10))

                if row_idx == 0:
                    # Header row: gray background + bold
                    add_cell_shading(cell, "D9D9D9")
                    set_cell_text_bold(cell)


def fix_table_spacing(doc):
    """Add spacing after tables so text doesn't crowd the bottom border."""
    body = doc.element.body
    children = list(body)
    count = 0
    for i, child in enumerate(children):
        if child.tag == qn('w:tbl') and i + 1 < len(children):
            next_el = children[i + 1]
            if next_el.tag == qn('w:p'):
                # Add space_before to the paragraph after the table
                pPr = next_el.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                    next_el.insert(0, pPr)
                existing = pPr.find(qn('w:spacing'))
                # Preserve existing after value, just set/update before
                before_twips = int(Pt(10) / 12700 * 20)
                if existing is not None:
                    existing.set(qn('w:before'), str(before_twips))
                else:
                    spacing = parse_xml(
                        f'<w:spacing {nsdecls("w")} w:before="{before_twips}"/>'
                    )
                    pPr.append(spacing)
                count += 1
    return count


def fix_code_blocks(doc):
    """Find and style code block paragraphs."""
    code_count = 0
    for paragraph in doc.paragraphs:
        if is_code_block(paragraph):
            style_code_block(paragraph)
            code_count += 1
    return code_count


def fix_heading_hierarchy(doc):
    """
    Normalize heading hierarchy to consistent styling:
    - H1: 20pt, black
    - H2: 16pt, black
    - H3: 14pt, #434343
    - H4: 12pt, #666666
    """
    heading_styles = {
        'Heading 1': {'size': Pt(20), 'color': RGBColor(0, 0, 0), 'bold': False},
        'Heading 2': {'size': Pt(16), 'color': RGBColor(0, 0, 0), 'bold': False},
        'Heading 3': {'size': Pt(14), 'color': RGBColor(0x43, 0x43, 0x43), 'bold': False},
        'Heading 4': {'size': Pt(12), 'color': RGBColor(0x66, 0x66, 0x66), 'bold': False},
        # h5/h6 in markdown - promote to h4 for better visibility
        'Heading 5': {'size': Pt(11), 'color': RGBColor(0x66, 0x66, 0x66), 'bold': True},
        'Heading 6': {'size': Pt(11), 'color': RGBColor(0x66, 0x66, 0x66), 'bold': True},
    }

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in heading_styles:
            config = heading_styles[style_name]
            for run in paragraph.runs:
                run.font.size = config['size']
                run.font.color.rgb = config['color']
                run.font.name = "Arial"
                if config.get('bold'):
                    run.bold = True


def set_paragraph_spacing(paragraph, before=None, after=None, line=None, line_rule=None):
    """Set spacing on a paragraph (before/after in Pt, line as multiplier)."""
    pPr = paragraph._p.get_or_add_pPr()

    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)

    attrs = f'{nsdecls("w")}'
    if before is not None:
        # Convert Pt to twips (1pt = 20 twips)
        attrs += f' w:before="{int(before / 12700 * 20)}"'
    if after is not None:
        attrs += f' w:after="{int(after / 12700 * 20)}"'
    if line is not None:
        # Line spacing in 240ths of a line (240 = single, 276 = 1.15, 360 = 1.5)
        attrs += f' w:line="{line}"'
        if line_rule:
            attrs += f' w:lineRule="{line_rule}"'

    spacing = parse_xml(f'<w:spacing {attrs}/>')
    pPr.append(spacing)


def is_list_item(paragraph):
    """Check if a paragraph is a list item (has numbering properties)."""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        return pPr.find(qn('w:numPr')) is not None
    return False


def fix_paragraph_spacing(doc):
    """Add spacing to paragraphs for better readability."""
    paragraphs = doc.paragraphs
    body_count = 0
    list_count = 0
    code_spacing_count = 0

    for i, p in enumerate(paragraphs):
        style_name = p.style.name if p.style else ""

        # Skip headings (they already have spacing from the template styles)
        if style_name.startswith('Heading') or style_name == 'Title' or style_name == 'Subtitle':
            continue

        # Code blocks: add spacing before first and after last in a group
        if is_code_block(p):
            prev_is_code = i > 0 and is_code_block(paragraphs[i - 1])
            next_is_code = i < len(paragraphs) - 1 and is_code_block(paragraphs[i + 1])

            before_val = Pt(8) if not prev_is_code else None
            after_val = Pt(8) if not next_is_code else Pt(0)

            set_paragraph_spacing(p, before=before_val, after=after_val)
            code_spacing_count += 1
            continue

        # List items: tight spacing within lists
        if is_list_item(p):
            set_paragraph_spacing(p, after=Pt(2), line=276)
            list_count += 1
            continue

        # Normal body text: comfortable spacing + 1.15 line height
        if style_name.lower() in ('normal', 'body text', 'first paragraph'):
            set_paragraph_spacing(p, after=Pt(8), line=276)
            body_count += 1

    return body_count, list_count, code_spacing_count


def postprocess(input_path, output_path):
    """Main post-processing function."""
    doc = Document(input_path)

    print(f"Processing: {input_path}")

    # 1. Fix tables
    table_count = len(doc.tables)
    fix_tables(doc)
    print(f"  Fixed {table_count} tables (borders + header shading)")

    # 2. Fix code blocks
    code_count = fix_code_blocks(doc)
    print(f"  Styled {code_count} code block paragraphs")

    # 3. Fix heading hierarchy
    fix_heading_hierarchy(doc)
    print(f"  Fixed heading styles")

    # 4. Fix paragraph spacing
    body_n, list_n, code_n = fix_paragraph_spacing(doc)
    print(f"  Fixed spacing: {body_n} body, {list_n} list items, {code_n} code blocks")

    # 5. Fix spacing after tables
    table_sp = fix_table_spacing(doc)
    print(f"  Added spacing after {table_sp} tables")

    # Save
    doc.save(output_path)
    print(f"  Saved to: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python postprocess_docx.py <input.docx> <output.docx>")
        sys.exit(1)

    postprocess(sys.argv[1], sys.argv[2])
